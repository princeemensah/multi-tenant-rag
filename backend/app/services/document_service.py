"""Document service with upload, processing, and retrieval logic."""
from __future__ import annotations

import logging
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import HTTPException, UploadFile, status
from sqlalchemy import delete
from sqlalchemy.orm import Session

from app.config import settings
from app.models.document import Document, DocumentChunk
from app.services.embedding_service import EmbeddingService
from app.services.vector_service import QdrantVectorService

try:  # Optional heavy dependencies used for rich text extraction
    import PyPDF2  # type: ignore
except Exception:  # pragma: no cover - safe fallback
    PyPDF2 = None

try:  # pragma: no cover - optional dependency
    from docx import Document as DocxDocument  # type: ignore
except Exception:  # pragma: no cover
    DocxDocument = None

logger = logging.getLogger(__name__)


class DocumentService:
    """Business logic for document ingestion and retrieval."""

    def __init__(self) -> None:
        self.upload_dir = Path(settings.upload_dir)
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024
        self.allowed_types = {ext.lower() for ext in settings.allowed_file_types}
        self.embedding_service = EmbeddingService()
        self.vector_service = QdrantVectorService()
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def upload_document(
        self,
        db: Session,
        tenant_id: str,
        file: UploadFile,
        metadata: Optional[Dict[str, Any]] = None,
        title: Optional[str] = None,
        tags: Optional[List[str]] = None,
    ) -> Document:
        self._validate_file(file)
        file_ext = self._infer_extension(file.filename)
        stored_name = f"{tenant_id}_{uuid.uuid4()}.{file_ext}"
        file_path = self.upload_dir / stored_name

        raw = await file.read()
        if len(raw) > self.max_file_size:
            raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="File exceeds size limit")

        try:
            file_path.write_bytes(raw)
        except Exception as exc:
            logger.error("Failed to persist uploaded file", extra={"error": str(exc)})
            raise HTTPException(status_code=500, detail="Failed to store uploaded file") from exc

        document = Document(
            tenant_id=tenant_id,
            filename=stored_name,
            original_filename=file.filename or stored_name,
            content_type=file.content_type or self._guess_mime(file_ext),
            file_size=len(raw),
            file_path=str(file_path),
            status="uploaded",
            title=title,
            tags=tags or [],
            doc_metadata=metadata or {},
        )

        db.add(document)
        db.commit()
        db.refresh(document)
        logger.info("Document uploaded", extra={"document_id": str(document.id), "tenant": tenant_id})
        return document

    def _validate_file(self, file: UploadFile) -> None:
        if not file.filename:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Filename missing")
        extension = self._infer_extension(file.filename)
        if extension not in self.allowed_types:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type")

    def _infer_extension(self, filename: Optional[str]) -> str:
        return (filename or "").rsplit(".", maxsplit=1)[-1].lower() if filename and "." in filename else "txt"

    def _guess_mime(self, extension: str) -> str:
        mapping = {
            "pdf": "application/pdf",
            "txt": "text/plain",
            "md": "text/markdown",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        return mapping.get(extension, "application/octet-stream")

    async def process_document(self, db: Session, document_id: str, tenant_id: str) -> bool:
        document = self.get_document(db, document_id, tenant_id)
        if not document:
            logger.error("Document not found", extra={"document_id": document_id, "tenant": tenant_id})
            return False

        if not await self.vector_service.init_collection():
            logger.error("Vector collection unavailable")
            return False

        document.status = "processing"
        document.processed_at = None
        db.commit()

        text = self._extract_text(document.file_path, document.content_type)
        if not text.strip():
            document.status = "failed"
            db.commit()
            logger.error("No text extracted", extra={"document_id": document_id})
            return False

        document.word_count = len(text.split())

        # Remove previous artefacts if reprocessing
        db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == document.id))
        await self.vector_service.delete_document(tenant_id, str(document.id))
        db.flush()

        chunks = self.embedding_service.chunk_text_for_embedding(text)
        if not chunks:
            document.status = "failed"
            db.commit()
            return False

        embedded_chunks = await self.embedding_service.embed_document_chunks(chunks)

        chunk_records: List[DocumentChunk] = []
        vector_payloads: List[Dict[str, Any]] = []

        for chunk in embedded_chunks:
            chunk_id = uuid.uuid4()
            vector_id = str(uuid.uuid4())
            start_val = chunk.get("start_char")
            end_val = chunk.get("end_char")
            page_number = chunk.get("page_number")
            record = DocumentChunk(
                id=chunk_id,
                document_id=document.id,
                tenant_id=document.tenant_id,
                chunk_index=int(chunk["chunk_index"]),
                text_content=str(chunk["text"]),
                chunk_size=int(chunk.get("chunk_size") or len(str(chunk["text"]))),
                start_char=int(start_val) if start_val is not None else None,
                end_char=int(end_val) if end_val is not None else None,
                page_number=int(page_number) if page_number is not None else None,
                vector_id=vector_id,
                embedding_model=str(chunk.get("embedding_model")),
                embedding_dimension=int(chunk.get("embedding_dimension") or 0),
                doc_metadata=document.doc_metadata or {},
            )
            chunk_records.append(record)

            vector_payloads.append(
                {
                    "document_id": str(document.id),
                    "chunk_id": str(chunk_id),
                    "text": record.text_content,
                    "embedding": chunk["embedding"],
                    "source": document.original_filename,
                    "page_number": record.page_number,
                    "chunk_index": record.chunk_index,
                    "tags": document.tags,
                    "metadata": {
                        "filename": document.original_filename,
                        "content_type": document.content_type,
                        "tags": document.tags,
                        "document_metadata": document.doc_metadata,
                        "start_char": record.start_char,
                        "end_char": record.end_char,
                    },
                }
            )

        db.add_all(chunk_records)
        db.commit()

        success = await self.vector_service.add_documents(tenant_id=str(document.tenant_id), documents=vector_payloads)
        if not success:
            document.status = "failed"
            db.commit()
            return False

        document.status = "processed"
        document.total_chunks = len(chunk_records)
        document.processed_chunks = len(chunk_records)
        document.collection_name = self.vector_service.default_collection
        document.embedding_model = self.embedding_service.model_name
        document.processed_at = datetime.utcnow()
        db.commit()
        logger.info("Document processed", extra={"document_id": str(document.id), "chunks": len(chunk_records)})
        return True

    async def delete_document(self, db: Session, document_id: str, tenant_id: str) -> bool:
        document = self.get_document(db, document_id, tenant_id)
        if not document:
            return False

        await self.vector_service.delete_document(tenant_id, str(document.id))
        db.delete(document)
        db.commit()

        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception as exc:  # pragma: no cover - cleanup best effort
            logger.warning("Failed to remove file", extra={"error": str(exc), "path": document.file_path})

        logger.info("Document deleted", extra={"document_id": document_id, "tenant": tenant_id})
        return True

    def get_document(self, db: Session, document_id: str, tenant_id: Optional[str] = None) -> Optional[Document]:
        query = db.query(Document).filter(Document.id == document_id)
        if tenant_id:
            query = query.filter(Document.tenant_id == tenant_id)
        return query.first()

    def list_documents(
        self,
        db: Session,
        tenant_id: str,
        skip: int = 0,
        limit: int = 100,
        status_filter: Optional[str] = None,
    ) -> List[Document]:
        query = db.query(Document).filter(Document.tenant_id == tenant_id)
        if status_filter:
            query = query.filter(Document.status == status_filter)
        return query.offset(skip).limit(limit).all()

    def update_document_status(self, db: Session, document_id: str, status_value: str) -> Document:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
        document.status = status_value
        db.commit()
        db.refresh(document)
        return document

    async def chunk_and_embed(self, text: str) -> List[Dict[str, Any]]:
        chunks = self.embedding_service.chunk_text_for_embedding(text)
        return await self.embedding_service.embed_document_chunks(chunks)

    def _extract_text(self, path: str, content_type: str) -> str:
        if content_type == "application/pdf" or path.lower().endswith(".pdf"):
            return self._extract_pdf(path)
        if (
            content_type.endswith("msword")
            or "wordprocessingml" in content_type
            or path.lower().endswith(".docx")
        ):
            return self._extract_docx(path)
        return self._extract_text_file(path)

    def _extract_pdf(self, path: str) -> str:
        if PyPDF2 is None:
            logger.warning("PyPDF2 unavailable; treating PDF as binary text")
            return self._extract_text_file(path)
        text_parts: List[str] = []
        with open(path, "rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            for index, page in enumerate(reader.pages):
                try:
                    extracted = page.extract_text() or ""
                except Exception:  # pragma: no cover - best effort
                    extracted = ""
                if extracted.strip():
                    text_parts.append(f"[Page {index + 1}]\n{extracted}")
        return "\n\n".join(text_parts)

    def _extract_docx(self, path: str) -> str:
        if DocxDocument is None:
            logger.warning("python-docx unavailable; treating DOCX as text")
            return self._extract_text_file(path)
        document = DocxDocument(path)
        lines = [para.text for para in document.paragraphs if para.text.strip()]
        return "\n\n".join(lines)

    def _extract_text_file(self, path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as handle:
            return handle.read()

